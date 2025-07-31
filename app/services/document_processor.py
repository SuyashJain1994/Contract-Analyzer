import io
import logging
from typing import Union
import PyPDF2
import docx
from pathlib import Path

from ..utils.exceptions import DocumentProcessingException

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing and extracting text from various document formats"""
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt'}
    
    async def extract_text(self, content: bytes, filename: str) -> str:
        """Extract text from document content based on file extension"""
        try:
            file_extension = Path(filename).suffix.lower()
            
            if file_extension not in self.supported_extensions:
                raise DocumentProcessingException(
                    f"Unsupported file format: {file_extension}",
                    status_code=400
                )
            
            if file_extension == '.pdf':
                return await self._extract_from_pdf(content)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_from_docx(content)
            elif file_extension == '.txt':
                return await self._extract_from_txt(content)
            else:
                raise DocumentProcessingException(
                    f"Handler not implemented for {file_extension}",
                    status_code=500
                )
                
        except DocumentProcessingException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise DocumentProcessingException(
                f"Failed to process document: {str(e)}",
                status_code=500
            )
    
    async def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                raise DocumentProcessingException(
                    "No text could be extracted from PDF. The document might be image-based or corrupted.",
                    status_code=400
                )
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise DocumentProcessingException(
                f"Failed to extract text from PDF: {str(e)}",
                status_code=500
            )
    
    async def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise DocumentProcessingException(
                    "No text could be extracted from document.",
                    status_code=400
                )
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise DocumentProcessingException(
                f"Failed to extract text from DOCX: {str(e)}",
                status_code=500
            )
    
    async def _extract_from_txt(self, content: bytes) -> str:
        """Extract text from TXT content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingException(
                "Could not decode text file with any supported encoding",
                status_code=400
            )
            
        except DocumentProcessingException:
            raise
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise DocumentProcessingException(
                f"Failed to extract text from TXT: {str(e)}",
                status_code=500
            )
    
    def validate_file(self, filename: str, content_length: int, max_size: int = 50 * 1024 * 1024) -> bool:
        """Validate file before processing"""
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise DocumentProcessingException(
                f"Unsupported file format: {file_extension}. Supported formats: {', '.join(self.supported_extensions)}",
                status_code=400
            )
        
        if content_length > max_size:
            raise DocumentProcessingException(
                f"File too large: {content_length} bytes. Maximum size: {max_size} bytes",
                status_code=413
            )
        
        return True